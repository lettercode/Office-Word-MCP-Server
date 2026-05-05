"""Tests for search_and_replace cross-run text matching (Bug 1)."""
import asyncio
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from word_document_server.tools.content_tools import search_and_replace
from word_document_server.utils.document_utils import find_and_replace_text


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_PKG_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _hyperlink_displays(path):
    """Return list of (rId, concatenated display text) for every w:hyperlink."""
    with zipfile.ZipFile(path) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    out = []
    for h in root.findall(f".//{{{W_NS}}}hyperlink"):
        rid = h.get(f"{{{R_NS}}}id")
        text = "".join(t.text or "" for t in h.findall(f".//{{{W_NS}}}t"))
        out.append((rid, text))
    return out


def _hyperlink_targets(path):
    """Return {rId: target} for every external hyperlink in the rels part."""
    with zipfile.ZipFile(path) as z:
        rels = ET.fromstring(z.read("word/_rels/document.xml.rels"))
    hyper_type = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    )
    return {
        r.get("Id"): r.get("Target")
        for r in rels.findall(f"{{{REL_PKG_NS}}}Relationship")
        if r.get("Type") == hyper_type
    }


class TestSingleRunReplace:
    """Regression: replacement within a single run still works."""

    def test_single_run_replace(self, make_docx):
        path = make_docx(paragraphs=["Hello World"])
        doc = Document(path)
        count = find_and_replace_text(doc, "Hello", "Goodbye")
        doc.save(path)
        doc2 = Document(path)
        assert doc2.paragraphs[0].text == "Goodbye World"
        assert count >= 1


class TestCrossRunReplace:
    """Text that spans multiple <w:r> elements must be matched and replaced."""

    def test_cross_run_replace(self, cross_run_docx):
        doc = Document(cross_run_docx)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(cross_run_docx)
        doc2 = Document(cross_run_docx)
        assert doc2.paragraphs[0].text == "Goodbye Earth"
        assert count >= 1

    def test_preserves_first_run_formatting(self, multi_run_formatted_docx):
        doc = Document(multi_run_formatted_docx)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(multi_run_formatted_docx)
        doc2 = Document(multi_run_formatted_docx)
        assert doc2.paragraphs[0].text == "Goodbye Earth"
        first_run = doc2.paragraphs[0].runs[0]
        assert first_run.bold is True
        assert count >= 1

    def test_table_cell_cross_run(self, table_docx):
        doc = Document(table_docx)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(table_docx)
        doc2 = Document(table_docx)
        cell_text = doc2.tables[0].cell(0, 0).text
        assert cell_text == "Goodbye Earth"
        assert count >= 1

    def test_no_match_returns_zero(self, make_docx):
        path = make_docx(paragraphs=["Hello World"])
        doc = Document(path)
        count = find_and_replace_text(doc, "Nonexistent", "Replacement")
        assert count == 0

    def test_multiple_occurrences(self, make_docx):
        path = make_docx(paragraphs=[
            {"runs": [{"text": "Hello "}, {"text": "World"}]},
            {"runs": [{"text": "Hello "}, {"text": "World"}]},
        ])
        doc = Document(path)
        count = find_and_replace_text(doc, "Hello World", "Goodbye Earth")
        doc.save(path)
        doc2 = Document(path)
        assert doc2.paragraphs[0].text == "Goodbye Earth"
        assert doc2.paragraphs[1].text == "Goodbye Earth"
        assert count == 2

    def test_toc_paragraphs_skipped(self, tmp_path):
        path = tmp_path / "toc_test.docx"
        doc = Document()
        p = doc.add_paragraph("Hello World")
        p.style = doc.styles.add_style("TOC 1", 1) if "TOC 1" not in [s.name for s in doc.styles] else doc.styles["TOC 1"]
        doc.save(str(path))
        doc2 = Document(str(path))
        count = find_and_replace_text(doc2, "Hello World", "Goodbye Earth")
        assert count == 0
        assert doc2.paragraphs[0].text == "Hello World"


class TestOutlinePrefixDiagnostic:
    """When find_text uses markdown heading prefix (e.g. '## Foo') but
    paragraph text is literal 'Foo' on a Heading style, return a hint
    rather than a silent 'No occurrences found'. See word-mcp-bug-report.md."""

    def test_markdown_heading_prefix_no_match_returns_hint(self, make_docx):
        path = make_docx(paragraphs=[
            "Intro text",
            {"style": "Heading 2", "runs": [{"text": "Performance"}]},
            "Body after heading",
        ])
        result = asyncio.run(search_and_replace(path, "## Performance", "## Performance (H1.1a)"))
        assert "No occurrences" in result
        assert "Performance" in result
        assert "Heading 2" in result
        # Paragraph index hint must point at the actual heading (index 1 here).
        assert "index 1" in result
        # Document must be unchanged on disk.
        doc2 = Document(path)
        assert doc2.paragraphs[1].text == "Performance"

    def test_markdown_heading_prefix_with_literal_match_still_replaces(self, make_docx):
        path = make_docx(paragraphs=["## Performance is literal here"])
        result = asyncio.run(search_and_replace(path, "## Performance", "## Speed"))
        assert "Replaced 1 occurrence" in result
        doc2 = Document(path)
        assert doc2.paragraphs[0].text == "## Speed is literal here"

    def test_wrong_prefix_level_returns_hint(self, make_docx):
        path = make_docx(paragraphs=[
            {"style": "Heading 1", "runs": [{"text": "Methods"}]},
        ])
        result = asyncio.run(search_and_replace(path, "### Methods", "### Methods updated"))
        assert "No occurrences" in result
        assert "Heading 1" in result
        # Level mismatch should be flagged so caller knows the outline hash count
        # did not match the actual heading level.
        assert "does not match" in result or "mismatch" in result.lower()

    def test_no_hint_when_no_heading_match(self, make_docx):
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "Performance"}]},
        ])
        result = asyncio.run(search_and_replace(path, "## Nonexistent", "irrelevant"))
        assert "No occurrences" in result
        assert "Hint" not in result
        assert "Heading" not in result

    def test_non_prefix_misses_are_unchanged(self, make_docx):
        path = make_docx(paragraphs=[
            {"style": "Heading 2", "runs": [{"text": "Performance"}]},
        ])
        result = asyncio.run(search_and_replace(path, "Nonexistent", "irrelevant"))
        assert result == "No occurrences of 'Nonexistent' found."


class TestNoHangWhenReplacementContainsSearchTerm:
    """Regression for the reported `search_and_replace` hang.

    The original `_replace_in_paragraph` re-tested `old_text in para.text`
    after every rewrite, so any `new_text` that itself contained `old_text`
    triggered an infinite loop with no error and no return. After the fix,
    each existing match is replaced exactly once.
    """

    def test_replacement_containing_search_term_does_not_hang(self, make_docx):
        path = make_docx(paragraphs=["Schedule. The team meets weekly."])
        doc = Document(path)
        count = find_and_replace_text(
            doc, "Schedule.", "Schedule.\n\nDetails follow."
        )
        doc.save(path)
        doc2 = Document(path)
        assert count == 1
        assert (
            doc2.paragraphs[0].text
            == "Schedule.\n\nDetails follow. The team meets weekly."
        )

    def test_replacement_idempotent_when_new_contains_old(self, make_docx):
        path = make_docx(paragraphs=["X marks the spot"])
        doc = Document(path)
        count = find_and_replace_text(doc, "X", "XX")
        doc.save(path)
        doc2 = Document(path)
        assert count == 1
        assert doc2.paragraphs[0].text == "XX marks the spot"

    def test_multiple_paragraph_matches_with_recursive_replacement(self, make_docx):
        path = make_docx(paragraphs=["foo one", "foo two"])
        doc = Document(path)
        count = find_and_replace_text(doc, "foo", "foo bar")
        doc.save(path)
        doc2 = Document(path)
        assert count == 2
        assert doc2.paragraphs[0].text == "foo bar one"
        assert doc2.paragraphs[1].text == "foo bar two"

    def test_two_occurrences_in_same_paragraph(self, make_docx):
        path = make_docx(paragraphs=["foo and foo"])
        doc = Document(path)
        count = find_and_replace_text(doc, "foo", "foo!")
        doc.save(path)
        doc2 = Document(path)
        assert count == 2
        assert doc2.paragraphs[0].text == "foo! and foo!"


class TestHyperlinkAwareReplace:
    """Bug B Symptom 1: search_and_replace must reach text inside w:hyperlink."""

    def test_replace_finds_all_hyperlink_and_plain_occurrences(self, hyperlink_docx):
        doc = Document(hyperlink_docx)
        count = find_and_replace_text(doc, "FOO", "QUX")
        doc.save(hyperlink_docx)
        # 5 occurrences total: para 0, para 1 hyperlink, para 2 hyperlink (cross-run),
        # para 3, table cell.
        assert count == 5

        doc2 = Document(hyperlink_docx)
        assert doc2.paragraphs[0].text == "QUX appears here"
        # Para 1: prefix run + hyperlink — combined visible text via paragraph.text
        # which in this python-docx version concatenates hyperlink runs too.
        assert "Visit " in doc2.paragraphs[1].text
        assert "QUX bar" in doc2.paragraphs[1].text
        assert "QUX baz" in doc2.paragraphs[2].text
        assert doc2.paragraphs[3].text == "Tail QUX"
        assert doc2.tables[0].cell(0, 0).text == "QUX in cell"

    def test_hyperlink_target_url_unchanged_after_replace(self, hyperlink_docx):
        targets_before = _hyperlink_targets(hyperlink_docx)
        doc = Document(hyperlink_docx)
        find_and_replace_text(doc, "FOO", "QUX")
        doc.save(hyperlink_docx)
        targets_after = _hyperlink_targets(hyperlink_docx)
        assert targets_before == targets_after  # rels untouched

    def test_hyperlink_displays_after_replace(self, hyperlink_docx):
        doc = Document(hyperlink_docx)
        find_and_replace_text(doc, "FOO", "QUX")
        doc.save(hyperlink_docx)
        displays = sorted(t for _, t in _hyperlink_displays(hyperlink_docx))
        # Both hyperlinks now show QUX display text. Cross-run hyperlink merged
        # into a single visible string but still a single hyperlink element.
        assert displays == ["QUX bar", "QUX baz"]

    def test_search_and_replace_message_reports_hyperlink_count(self, hyperlink_docx):
        result = asyncio.run(search_and_replace(hyperlink_docx, "FOO", "QUX"))
        # Must report the total accurately and disclose that some matches
        # lived inside hyperlink display text.
        assert "Replaced 5 occurrence" in result
        assert "hyperlink" in result.lower()

    def test_document_reopens_cleanly_after_hyperlink_replace(self, hyperlink_docx):
        doc = Document(hyperlink_docx)
        find_and_replace_text(doc, "FOO", "QUX")
        doc.save(hyperlink_docx)
        # Re-opening must not raise (catches XML structural breakage).
        Document(hyperlink_docx)


class TestEmptyHyperlinkCleanup:
    """Bug B Symptom 2: replacing the *entire* display text of a hyperlink
    with empty must NOT leave a zombie ``<w:hyperlink>`` element behind."""

    def test_clearing_full_hyperlink_text_removes_element(self, hyperlink_docx):
        doc = Document(hyperlink_docx)
        # Replace just the display label of the single-run hyperlink with "".
        find_and_replace_text(doc, "FOO bar", "")
        doc.save(hyperlink_docx)

        displays = _hyperlink_displays(hyperlink_docx)
        # The "FOO bar" hyperlink must be gone; the cross-run "FOO baz" remains.
        assert ("FOO baz" in [t for _, t in displays])
        assert not any(t == "" for _, t in displays)
        assert all("FOO bar" not in t for _, t in displays)

    def test_partial_hyperlink_clear_keeps_hyperlink(self, hyperlink_docx):
        doc = Document(hyperlink_docx)
        # 'FOO ' inside 'FOO bar' — leaves 'bar' behind. Hyperlink must remain.
        find_and_replace_text(doc, "FOO ", "")
        doc.save(hyperlink_docx)

        displays = [t for _, t in _hyperlink_displays(hyperlink_docx)]
        # Some hyperlink still has 'bar' as its (now-shorter) display text.
        assert any("bar" in t for t in displays)
        # No hyperlink lost its element entirely.
        assert "" not in displays

    def test_empty_hyperlink_element_not_present_in_xml(self, hyperlink_docx):
        doc = Document(hyperlink_docx)
        find_and_replace_text(doc, "FOO bar", "")
        doc.save(hyperlink_docx)

        with zipfile.ZipFile(hyperlink_docx) as z:
            root = ET.fromstring(z.read("word/document.xml"))
        # No hyperlink element should have an empty visible text.
        for h in root.findall(f".//{{{W_NS}}}hyperlink"):
            visible = "".join(t.text or "" for t in h.findall(f".//{{{W_NS}}}t"))
            assert visible != "", "empty <w:hyperlink> survived cleanup"
