import pytest
from docx import Document
from docx.shared import Pt, RGBColor


@pytest.fixture
def make_docx(tmp_path):
    """Factory fixture: creates a .docx with specified paragraph structure."""
    def _make(filename="test.docx", paragraphs=None):
        path = tmp_path / filename
        doc = Document()
        for p_spec in (paragraphs or []):
            if isinstance(p_spec, str):
                doc.add_paragraph(p_spec)
            elif isinstance(p_spec, dict):
                style = p_spec.get("style", "Normal")
                para = doc.add_paragraph("", style=style)
                for run_spec in p_spec.get("runs", []):
                    run = para.add_run(run_spec["text"])
                    if "bold" in run_spec:
                        run.bold = run_spec["bold"]
                    if "italic" in run_spec:
                        run.italic = run_spec["italic"]
                    if "font_size" in run_spec:
                        run.font.size = Pt(run_spec["font_size"])
                    if "font_name" in run_spec:
                        run.font.name = run_spec["font_name"]
        doc.save(str(path))
        return str(path)
    return _make


@pytest.fixture
def cross_run_docx(make_docx):
    """'Hello World' split across two runs."""
    return make_docx(paragraphs=[
        {"runs": [{"text": "Hello "}, {"text": "World"}]},
        "Simple paragraph",
    ])


@pytest.fixture
def multi_run_formatted_docx(make_docx):
    """'Hello World' split across runs with different formatting."""
    return make_docx(paragraphs=[
        {"runs": [
            {"text": "Hello ", "bold": True, "font_size": 12},
            {"text": "World", "bold": False, "font_size": 14},
        ]},
    ])


@pytest.fixture
def heading_docx(make_docx):
    """Document with headings and content blocks."""
    return make_docx(paragraphs=[
        {"style": "Heading 1", "runs": [{"text": "Section One"}]},
        "Content under section one.",
        "More content.",
        {"style": "Heading 1", "runs": [{"text": "Section Two"}]},
        "Content under section two.",
    ])


@pytest.fixture
def table_docx(tmp_path):
    """Table with cross-run text in cell(0,0)."""
    path = tmp_path / "table_test.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell.text = ""
    para = cell.paragraphs[0]
    para.add_run("Hello ")
    para.add_run("World")
    table.cell(0, 1).text = "Other cell"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def anchor_docx(tmp_path):
    """START/END anchor paragraphs with content between."""
    path = tmp_path / "anchor_test.docx"
    doc = Document()
    doc.add_paragraph("--- START ANCHOR ---")
    doc.add_paragraph("Content to replace 1")
    doc.add_paragraph("Content to replace 2")
    doc.add_paragraph("--- END ANCHOR ---")
    doc.add_paragraph("After the anchors")
    doc.save(str(path))
    return str(path)


@pytest.fixture
def hyperlink_docx(tmp_path):
    """Document covering every place hyperlink-embedded text can appear.

    Layout (paragraph index → contents):
      0: plain paragraph "FOO appears here"
      1: paragraph "Visit " + hyperlink "FOO bar" (single run) → https://example.com/foo
      2: paragraph "Cross-run " + hyperlink whose display is split across
         two runs, "FOO" and " baz" → https://example.com/foobaz
      3: paragraph "Tail FOO" (plain trailing match)
      table[0,0]: "FOO in cell"
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    path = tmp_path / "hyperlink_test.docx"
    doc = Document()

    # Para 0: plain text only
    doc.add_paragraph("FOO appears here")

    # Para 1: prefix + hyperlink with single display run
    p1 = doc.add_paragraph("Visit ")
    from word_document_server.core.hyperlinks import add_hyperlink_run
    add_hyperlink_run(p1, "https://example.com/foo", "FOO bar")

    # Para 2: prefix + hyperlink whose display is split into 2 runs
    p2 = doc.add_paragraph("Cross-run ")
    part = p2.part
    r_id = part.relate_to("https://example.com/foobaz", RT.HYPERLINK, is_external=True)
    hyper = OxmlElement("w:hyperlink")
    hyper.set(qn("r:id"), r_id)
    for chunk in ("FOO", " baz"):
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = chunk
        r.append(t)
        hyper.append(r)
    p2._p.append(hyper)

    # Para 3: plain trailing match
    doc.add_paragraph("Tail FOO")

    # Table cell with FOO
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "FOO in cell"

    doc.save(str(path))
    return str(path)


@pytest.fixture
def nbsp_anchor_docx(tmp_path):
    """Anchors with NBSP and ZWSP."""
    path = tmp_path / "nbsp_anchor_test.docx"
    doc = Document()
    doc.add_paragraph("---\u00a0START ANCHOR\u00a0---")
    doc.add_paragraph("Content to replace")
    doc.add_paragraph("---\u200bEND ANCHOR\u200b---")
    doc.save(str(path))
    return str(path)
