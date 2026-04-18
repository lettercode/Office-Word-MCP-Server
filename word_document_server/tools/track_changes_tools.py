"""Track changes tools for Word Document Server.

Provides tools for creating and managing tracked changes (insertions, deletions,
replacements) in Word documents using the docx-editor library.
"""
import os
import json
import getpass
from typing import Optional

from docx_editor import Document as DocxEditorDocument
from docx_editor.exceptions import TextNotFoundError

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable
from word_document_server.utils.docx_zip_utils import strip_meta_json
from word_document_server.utils.anchor_utils import normalize_paragraph_runs_for_anchor


def _get_author(author: Optional[str]) -> str:
    """Get the author name, defaulting to the system username."""
    if author and author.strip():
        return author.strip()
    return getpass.getuser()


def _open_tracked_document(filename: str, author: str) -> DocxEditorDocument:
    """Open a document for tracked changes editing."""
    return DocxEditorDocument.open(filename, author=author, force_recreate=True)


def _save_and_sanitize(doc: DocxEditorDocument, filename: str) -> None:
    """Save via docx-editor, then strip its stray meta.json from the zip.

    docx-editor writes its workspace meta.json inside the unpacked OPC tree
    and zips it into the .docx, which Word flags as "unreadable content".
    """
    doc.save()
    strip_meta_json(filename)


async def replace_with_track_changes(
    filename: str,
    find_text: str,
    replace_text: str,
    author: str = None,
    occurrence: int = None,
) -> str:
    """Replace text with tracked changes.

    Args:
        filename: Path to the Word document
        find_text: Text to find and replace
        replace_text: Replacement text
        author: Author name for the tracked change (defaults to system username)
        occurrence: Which occurrence to replace (0-indexed). None = replace all.

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    if not find_text:
        return json.dumps({"success": False, "error": "find_text cannot be empty"})

    author = _get_author(author)
    # Split runs so that a cross-run find_text survives docx-editor's
    # single-run matcher. If the text is truly absent, fall through to the
    # existing "no matches" response.
    normalize_paragraph_runs_for_anchor(filename, find_text)
    doc = None
    try:
        doc = _open_tracked_document(filename, author)
        match_count = doc.count_matches(find_text)
        if match_count == 0:
            return json.dumps({"success": True, "message": f"No matches found for '{find_text}'", "replacements": 0})

        if occurrence is not None:
            if occurrence < 0 or occurrence >= match_count:
                return json.dumps({"success": False, "error": f"Occurrence {occurrence} out of range. Found {match_count} matches (0-indexed)."})
            doc.replace(find_text, replace_text, occurrence=occurrence)
            replaced = 1
        else:
            # Replace all occurrences in reverse order to preserve positions
            for i in range(match_count - 1, -1, -1):
                doc.replace(find_text, replace_text, occurrence=i)
            replaced = match_count

        _save_and_sanitize(doc, filename)
        return json.dumps({"success": True, "replacements": replaced, "author": author})
    except TextNotFoundError:
        return json.dumps({"success": True, "message": f"No matches found for '{find_text}'", "replacements": 0})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to replace with track changes: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def delete_with_track_changes(
    filename: str,
    text: str,
    author: str = None,
    occurrence: int = None,
) -> str:
    """Mark text as deleted with tracked changes.

    Args:
        filename: Path to the Word document
        text: Text to mark as deleted
        author: Author name for the tracked change (defaults to system username)
        occurrence: Which occurrence to delete (0-indexed). None = delete all.

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    if not text:
        return json.dumps({"success": False, "error": "text cannot be empty"})

    author = _get_author(author)
    doc = None
    try:
        doc = _open_tracked_document(filename, author)
        match_count = doc.count_matches(text)
        if match_count == 0:
            return json.dumps({"success": True, "message": f"No matches found for '{text}'", "deletions": 0})

        if occurrence is not None:
            if occurrence < 0 or occurrence >= match_count:
                return json.dumps({"success": False, "error": f"Occurrence {occurrence} out of range. Found {match_count} matches (0-indexed)."})
            doc.delete(text, occurrence=occurrence)
            deleted = 1
        else:
            for i in range(match_count - 1, -1, -1):
                doc.delete(text, occurrence=i)
            deleted = match_count

        _save_and_sanitize(doc, filename)
        return json.dumps({"success": True, "deletions": deleted, "author": author})
    except TextNotFoundError:
        return json.dumps({"success": True, "message": f"No matches found for '{text}'", "deletions": 0})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to delete with track changes: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def insert_after_with_track_changes(
    filename: str,
    anchor_text: str,
    text_to_insert: str,
    author: str = None,
    occurrence: int = 0,
) -> str:
    """Insert text after anchor with tracked changes.

    Args:
        filename: Path to the Word document
        anchor_text: Text to find as insertion point
        text_to_insert: Text to insert after the anchor
        author: Author name for the tracked change (defaults to system username)
        occurrence: Which occurrence of anchor to use (0-indexed, default 0)

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    if not anchor_text:
        return json.dumps({"success": False, "error": "anchor_text cannot be empty"})

    author = _get_author(author)
    normalize_paragraph_runs_for_anchor(filename, anchor_text)
    doc = None
    try:
        doc = _open_tracked_document(filename, author)
        doc.insert_after(anchor_text, text_to_insert, occurrence=occurrence)
        _save_and_sanitize(doc, filename)
        return json.dumps({"success": True, "inserted": text_to_insert, "after": anchor_text, "author": author})
    except TextNotFoundError:
        return json.dumps({"success": True, "message": f"Anchor text '{anchor_text}' not found", "inserted": False})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to insert after with track changes: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def insert_before_with_track_changes(
    filename: str,
    anchor_text: str,
    text_to_insert: str,
    author: str = None,
    occurrence: int = 0,
) -> str:
    """Insert text before anchor with tracked changes.

    Args:
        filename: Path to the Word document
        anchor_text: Text to find as insertion point
        text_to_insert: Text to insert before the anchor
        author: Author name for the tracked change (defaults to system username)
        occurrence: Which occurrence of anchor to use (0-indexed, default 0)

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    if not anchor_text:
        return json.dumps({"success": False, "error": "anchor_text cannot be empty"})

    author = _get_author(author)
    normalize_paragraph_runs_for_anchor(filename, anchor_text)
    doc = None
    try:
        doc = _open_tracked_document(filename, author)
        doc.insert_before(anchor_text, text_to_insert, occurrence=occurrence)
        _save_and_sanitize(doc, filename)
        return json.dumps({"success": True, "inserted": text_to_insert, "before": anchor_text, "author": author})
    except TextNotFoundError:
        return json.dumps({"success": True, "message": f"Anchor text '{anchor_text}' not found", "inserted": False})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to insert before with track changes: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def list_revisions(
    filename: str,
    author: str = None,
) -> str:
    """List all tracked changes in a document.

    Args:
        filename: Path to the Word document
        author: If provided, filter revisions by this author

    Returns:
        JSON string with list of revisions
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    doc = None
    try:
        doc = _open_tracked_document(filename, _get_author(author))
        revisions = doc.list_revisions(author=author)
        revisions_data = [
            {
                "id": r.id,
                "type": r.type,
                "author": r.author,
                "date": r.date.isoformat() if r.date else None,
                "text": r.text,
            }
            for r in revisions
        ]
        return json.dumps({"success": True, "revisions": revisions_data, "total": len(revisions_data)})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to list revisions: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def accept_revision(
    filename: str,
    revision_id: int,
) -> str:
    """Accept a single tracked change by ID.

    Args:
        filename: Path to the Word document
        revision_id: ID of the revision to accept

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    doc = None
    try:
        doc = _open_tracked_document(filename, _get_author(None))
        result = doc.accept_revision(revision_id)
        if result:
            _save_and_sanitize(doc, filename)
            return json.dumps({"success": True, "accepted": revision_id})
        else:
            return json.dumps({"success": False, "error": f"Revision {revision_id} not found"})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to accept revision: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def reject_revision(
    filename: str,
    revision_id: int,
) -> str:
    """Reject a single tracked change by ID.

    Args:
        filename: Path to the Word document
        revision_id: ID of the revision to reject

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    doc = None
    try:
        doc = _open_tracked_document(filename, _get_author(None))
        result = doc.reject_revision(revision_id)
        if result:
            _save_and_sanitize(doc, filename)
            return json.dumps({"success": True, "rejected": revision_id})
        else:
            return json.dumps({"success": False, "error": f"Revision {revision_id} not found"})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to reject revision: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def accept_all_revisions(
    filename: str,
    author: str = None,
) -> str:
    """Accept all tracked changes, optionally filtered by author.

    Args:
        filename: Path to the Word document
        author: If provided, only accept revisions by this author

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    doc = None
    try:
        doc = _open_tracked_document(filename, _get_author(None))
        count = doc.accept_all(author=author)
        _save_and_sanitize(doc, filename)
        return json.dumps({"success": True, "accepted": count, "author_filter": author})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to accept all revisions: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def reject_all_revisions(
    filename: str,
    author: str = None,
) -> str:
    """Reject all tracked changes, optionally filtered by author.

    Args:
        filename: Path to the Word document
        author: If provided, only reject revisions by this author

    Returns:
        JSON string with result
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})

    doc = None
    try:
        doc = _open_tracked_document(filename, _get_author(None))
        count = doc.reject_all(author=author)
        _save_and_sanitize(doc, filename)
        return json.dumps({"success": True, "rejected": count, "author_filter": author})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to reject all revisions: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def get_visible_text(
    filename: str,
) -> str:
    """Get the visible text of a document (insertions included, deletions excluded).

    Args:
        filename: Path to the Word document

    Returns:
        JSON string with the visible text
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    doc = None
    try:
        doc = _open_tracked_document(filename, _get_author(None))
        text = doc.get_visible_text()
        return json.dumps({"success": True, "text": text})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to get visible text: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


async def count_tracked_matches(
    filename: str,
    text: str,
) -> str:
    """Count occurrences of text in the visible document content.

    Counts in the visible text (insertions included, deletions excluded).

    Args:
        filename: Path to the Word document
        text: Text to count occurrences of

    Returns:
        JSON string with match count
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    if not text:
        return json.dumps({"success": False, "error": "text cannot be empty"})

    doc = None
    try:
        doc = _open_tracked_document(filename, _get_author(None))
        count = doc.count_matches(text)
        return json.dumps({"success": True, "text": text, "count": count})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to count matches: {str(e)}"})
    finally:
        if doc:
            try:
                doc.close(cleanup=True)
            except Exception:
                pass


# ── Bug 5: tracked variants for plain content-insertion tools ────────────

from datetime import datetime, timezone
from docx import Document as PyDocxDocument
from docx.oxml.ns import qn
from lxml import etree


def _next_ins_id(body) -> int:
    """Pick an unused w:id for a new <w:ins> by scanning existing w:ins/w:del."""
    highest = -1
    for el in body.iter():
        if el.tag in (qn("w:ins"), qn("w:del")):
            raw = el.get(qn("w:id"))
            if raw is not None:
                try:
                    highest = max(highest, int(raw))
                except ValueError:
                    pass
    return highest + 1


def _build_tracked_paragraph(text: str, author: str, style_id=None):
    """Construct a <w:p> whose run is wrapped in <w:ins>.

    w:id on the w:ins is a placeholder (0); caller must set a body-unique
    value before inserting into the document."""
    _tmp_root = etree.Element(qn("w:root"))
    p = etree.SubElement(_tmp_root, qn("w:p"))
    _tmp_root.remove(p)
    if style_id:
        pPr = etree.SubElement(p, qn("w:pPr"))
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
        pStyle.set(qn("w:val"), style_id)
    ins = etree.SubElement(p, qn("w:ins"))
    ins.set(qn("w:id"), "0")
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    r = etree.SubElement(ins, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set(qn("xml:space"), "preserve")
    t.text = text
    return p


def _resolve_heading_style_id(doc, level: int) -> str:
    target_name = f"Heading {level}"
    for style in doc.styles:
        if style.name == target_name:
            return style.style_id
    return f"Heading{level}"


async def add_paragraph_with_track_changes(
    filename: str,
    text: str,
    style: Optional[str] = None,
    author: str = None,
) -> str:
    """Append a paragraph wrapped in <w:ins> so it appears as a tracked insertion."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})
    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})
    if not text:
        return json.dumps({"success": False, "error": "text cannot be empty"})

    author = _get_author(author)
    try:
        doc = PyDocxDocument(filename)
        body = doc.element.body
        style_id = None
        if style:
            for s in doc.styles:
                if s.name == style:
                    style_id = s.style_id
                    break
            if style_id is None:
                style_id = style
        new_p = _build_tracked_paragraph(text, author, style_id=style_id)
        new_p.find(qn("w:ins")).set(qn("w:id"), str(_next_ins_id(body)))
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(new_p)
        else:
            body.append(new_p)
        doc.save(filename)
        return json.dumps({"success": True, "text": text, "author": author, "style": style})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to add tracked paragraph: {str(e)}"})


async def add_heading_with_track_changes(
    filename: str,
    text: str,
    level: int = 1,
    author: str = None,
) -> str:
    """Append a heading wrapped in <w:ins>. level is 1-9 (maps to Heading N)."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})
    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})
    if not text:
        return json.dumps({"success": False, "error": "text cannot be empty"})
    if not (1 <= level <= 9):
        return json.dumps({"success": False, "error": "level must be between 1 and 9"})

    author = _get_author(author)
    try:
        doc = PyDocxDocument(filename)
        body = doc.element.body
        style_id = _resolve_heading_style_id(doc, level)
        new_p = _build_tracked_paragraph(text, author, style_id=style_id)
        new_p.find(qn("w:ins")).set(qn("w:id"), str(_next_ins_id(body)))
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(new_p)
        else:
            body.append(new_p)
        doc.save(filename)
        return json.dumps({"success": True, "text": text, "level": level, "author": author})
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to add tracked heading: {str(e)}"})


async def insert_line_or_paragraph_near_text_with_track_changes(
    filename: str,
    anchor_text: str,
    text_to_insert: str,
    position: str = "after",
    author: str = None,
) -> str:
    """Insert a new tracked paragraph before/after the paragraph whose text
    contains anchor_text. Anchor matching is paragraph-text (cross-run)."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})
    writeable, error = check_file_writeable(filename)
    if not writeable:
        return json.dumps({"success": False, "error": error})
    if not anchor_text:
        return json.dumps({"success": False, "error": "anchor_text cannot be empty"})
    if position not in ("after", "before"):
        return json.dumps({"success": False, "error": "position must be 'after' or 'before'"})

    author = _get_author(author)
    try:
        doc = PyDocxDocument(filename)
        body = doc.element.body
        anchor_p = None
        for p in doc.paragraphs:
            if anchor_text in p.text:
                anchor_p = p._element
                break
        if anchor_p is None:
            return json.dumps({"success": False, "error": f"Anchor text '{anchor_text}' not found in document"})

        new_p = _build_tracked_paragraph(text_to_insert, author)
        new_p.find(qn("w:ins")).set(qn("w:id"), str(_next_ins_id(body)))
        if position == "after":
            anchor_p.addnext(new_p)
        else:
            anchor_p.addprevious(new_p)
        doc.save(filename)
        return json.dumps({
            "success": True,
            "inserted": text_to_insert,
            "position": position,
            "anchor_text": anchor_text,
            "author": author,
        })
    except Exception as e:
        return json.dumps({"success": False, "error": f"Failed to insert tracked line: {str(e)}"})
