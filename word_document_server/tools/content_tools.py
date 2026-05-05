"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
import re
from typing import List, Optional, Dict, Any
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import find_and_replace_text, find_and_replace_text_detailed, insert_header_near_text, insert_numbered_list_near_text, insert_line_or_paragraph_near_text, replace_paragraph_block_below_header, replace_block_between_manual_anchors, replace_paragraph_text, replace_paragraph_range, delete_paragraph_range, diagnose_outline_prefix_miss
from word_document_server.core.styles import ensure_heading_style, ensure_table_style
from word_document_server.core.hyperlinks import add_hyperlink_run, wrap_run_as_hyperlink


async def add_heading(filename: str, text: str, level: int = 1,
                      font_name: Optional[str] = None, font_size: Optional[int] = None,
                      bold: Optional[bool] = None, italic: Optional[bool] = None,
                      border_bottom: bool = False) -> str:
    """Add a heading to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
        font_name: Font family (e.g., 'Helvetica')
        font_size: Font size in points (e.g., 14)
        bold: True/False for bold text
        italic: True/False for italic text
        border_bottom: True to add bottom border (for section headers)
    """
    filename = ensure_docx_extension(filename)

    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"

    # Validate level range
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        # Ensure heading styles exist
        ensure_heading_style(doc)

        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            heading = doc.add_paragraph(text)
            heading.style = doc.styles['Normal']
            if heading.runs:
                run = heading.runs[0]
                run.bold = True
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)

        # Apply formatting to all runs in the heading
        if any([font_name, font_size, bold is not None, italic is not None]):
            for run in heading.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic

        # Add bottom border if requested
        if border_bottom:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            pPr = heading._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')  # 0.5pt border
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')

            pBdr.append(bottom)
            pPr.append(pBdr)

        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


async def add_paragraph(filename: str, text: str, style: Optional[str] = None,
                        font_name: Optional[str] = None, font_size: Optional[int] = None,
                        bold: Optional[bool] = None, italic: Optional[bool] = None,
                        color: Optional[str] = None) -> str:
    """Add a paragraph to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
        font_name: Font family (e.g., 'Helvetica', 'Times New Roman')
        font_size: Font size in points (e.g., 14, 36)
        bold: True/False for bold text
        italic: True/False for italic text
        color: RGB color as hex string (e.g., '000000' for black)
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph(text)

        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"

        # Apply formatting to all runs in the paragraph
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    # Remove any '#' prefix if present
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)

        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


_MD_LINK_PATTERN = re.compile(r'\[([^\]]+)\]\(([^)\s]+)\)')


async def add_hyperlink(filename: str, url: str, text: Optional[str] = None,
                        style: Optional[str] = None,
                        font_name: Optional[str] = None, font_size: Optional[int] = None,
                        bold: Optional[bool] = None, italic: Optional[bool] = None) -> str:
    """Append a new paragraph containing a single clickable hyperlink.

    Args:
        filename: Path to the Word document.
        url: Destination URL. Scheme is added if missing.
        text: Visible link text. Defaults to the URL.
        style: Paragraph style name (optional). The link run always uses the
            "Hyperlink" character style when the document has it.
        font_name/font_size/bold/italic: Optional run-level overrides.
    """
    filename = ensure_docx_extension(filename)

    if not url or not isinstance(url, str):
        return "Invalid parameter: url must be a non-empty string"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph()
        if style:
            try:
                paragraph.style = style
            except KeyError:
                paragraph.style = doc.styles['Normal']

        add_hyperlink_run(
            paragraph, url, text or url,
            bold=bold, italic=italic,
            font_name=font_name, font_size=font_size,
        )

        doc.save(filename)
        return f"Hyperlink added to {filename}"
    except Exception as e:
        return f"Failed to add hyperlink: {str(e)}"


async def add_paragraph_with_hyperlinks(filename: str,
                                        segments: List[Dict[str, Any]],
                                        style: Optional[str] = None) -> str:
    """Append a paragraph composed of mixed plain and hyperlink segments.

    Args:
        filename: Path to the Word document.
        segments: Ordered list of segment dicts. Each segment must include
            ``text``. Segments with a non-empty ``url`` become hyperlinks;
            others become plain runs. Optional per-segment keys:
            ``bold``, ``italic``, ``font_name``, ``font_size``, ``color``.
        style: Paragraph style name (optional).
    """
    filename = ensure_docx_extension(filename)

    if not isinstance(segments, list) or not segments:
        return "Invalid parameter: segments must be a non-empty list"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph()
        if style:
            try:
                paragraph.style = style
            except KeyError:
                paragraph.style = doc.styles['Normal']

        link_count = 0
        for seg in segments:
            if not isinstance(seg, dict):
                return "Invalid parameter: each segment must be an object"
            text = seg.get("text", "")
            if text is None:
                text = ""
            url = seg.get("url")
            if url:
                add_hyperlink_run(
                    paragraph, url, text or url,
                    bold=seg.get("bold"),
                    italic=seg.get("italic"),
                    font_name=seg.get("font_name"),
                    font_size=seg.get("font_size"),
                )
                link_count += 1
            else:
                run = paragraph.add_run(text)
                if seg.get("bold") is not None:
                    run.bold = seg["bold"]
                if seg.get("italic") is not None:
                    run.italic = seg["italic"]
                if seg.get("font_name"):
                    run.font.name = seg["font_name"]
                if seg.get("font_size"):
                    run.font.size = Pt(seg["font_size"])
                if seg.get("color"):
                    run.font.color.rgb = RGBColor.from_string(seg["color"].lstrip("#"))

        doc.save(filename)
        return f"Paragraph with {link_count} hyperlink(s) added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph with hyperlinks: {str(e)}"


def _iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _convert_markdown_links_in_paragraph(paragraph) -> int:
    """Rewrite any `[text](url)` occurrences in paragraph.text as real hyperlinks.

    Collapses the paragraph into a single concatenated string, then rebuilds
    the paragraph's run/hyperlink children in order. Existing non-link run
    formatting is lost for segments that get rewritten — acceptable for a
    v1 conversion tool. Paragraphs without any markdown link are untouched.
    """
    text = paragraph.text
    if not _MD_LINK_PATTERN.search(text):
        return 0

    segments = []
    last_end = 0
    for m in _MD_LINK_PATTERN.finditer(text):
        if m.start() > last_end:
            segments.append(("text", text[last_end:m.start()]))
        segments.append(("link", m.group(1), m.group(2)))
        last_end = m.end()
    if last_end < len(text):
        segments.append(("text", text[last_end:]))

    p_elem = paragraph._p
    for child in list(p_elem):
        tag = child.tag
        if tag == qn("w:r") or tag == qn("w:hyperlink"):
            p_elem.remove(child)

    link_count = 0
    for seg in segments:
        if seg[0] == "text":
            if seg[1]:
                paragraph.add_run(seg[1])
        else:
            _, link_text, url = seg
            add_hyperlink_run(paragraph, url, link_text)
            link_count += 1
    return link_count


async def convert_markdown_links(filename: str) -> str:
    """Convert every `[text](url)` in the document into a real hyperlink.

    Scans body paragraphs and table-cell paragraphs. Returns a summary of
    how many links were converted.
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        total = 0
        for para in _iter_all_paragraphs(doc):
            if para.style and para.style.name.startswith("TOC"):
                continue
            total += _convert_markdown_links_in_paragraph(para)

        if total == 0:
            return f"No markdown links found in {filename}"
        doc.save(filename)
        return f"Converted {total} markdown link(s) to hyperlinks in {filename}"
    except Exception as e:
        return f"Failed to convert markdown links: {str(e)}"


async def convert_text_to_hyperlink(filename: str, target_text: str, url: str,
                                    occurrence: int = 1) -> str:
    """Turn an existing text span in the document into a hyperlink.

    For v1 simplicity, operates on paragraphs where ``target_text`` appears
    intact in ``paragraph.text``. If the span is split across multiple
    runs, the paragraph is first flattened to a single run. Surrounding
    text is preserved; formatting on the target span is replaced by the
    Hyperlink style (or blue/underline fallback).

    Args:
        filename: Path to the Word document.
        target_text: Text span to convert.
        url: Target URL.
        occurrence: 1-based occurrence to replace (default 1). Use 0 to
            convert every occurrence.
    """
    filename = ensure_docx_extension(filename)

    if not target_text:
        return "Invalid parameter: target_text must be non-empty"
    if not url:
        return "Invalid parameter: url must be non-empty"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        remaining = None if occurrence == 0 else max(1, int(occurrence))
        seen = 0
        replaced = 0

        for para in _iter_all_paragraphs(doc):
            if para.style and para.style.name.startswith("TOC"):
                continue
            text = para.text
            if target_text not in text:
                continue

            # Count and decide which hits to replace in this paragraph.
            indices = []
            start = 0
            while True:
                idx = text.find(target_text, start)
                if idx < 0:
                    break
                seen += 1
                if remaining is None or seen == (occurrence if occurrence > 0 else seen):
                    indices.append(idx)
                    if remaining is not None:
                        break
                start = idx + len(target_text)

            if not indices:
                continue

            # Flatten paragraph into before/target/after segments in order.
            # Build by scanning text positions; only the first matching index
            # per iteration is handled (re-scan after each replacement).
            p_elem = para._p
            # Capture plain text split into chunks around matches.
            chunks = []
            cursor = 0
            for idx in indices:
                if idx > cursor:
                    chunks.append(("text", text[cursor:idx]))
                chunks.append(("link", target_text))
                cursor = idx + len(target_text)
            if cursor < len(text):
                chunks.append(("text", text[cursor:]))

            for child in list(p_elem):
                tag = child.tag
                if tag == qn("w:r") or tag == qn("w:hyperlink"):
                    p_elem.remove(child)

            for kind, chunk_text in chunks:
                if kind == "text":
                    if chunk_text:
                        para.add_run(chunk_text)
                else:
                    add_hyperlink_run(para, url, chunk_text)
                    replaced += 1

            if remaining is not None and replaced >= remaining:
                break

        if replaced == 0:
            return f"Target text not found in {filename}"
        doc.save(filename)
        return f"Converted {replaced} occurrence(s) of target text to hyperlink in {filename}"
    except Exception as e:
        return f"Failed to convert text to hyperlink: {str(e)}"


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def add_table_of_contents(filename: str, title: str = "Table of Contents", max_level: int = 3) -> str:
    """Add a table of contents to a Word document based on heading styles.
    
    Args:
        filename: Path to the Word document
        title: Optional title for the table of contents
        max_level: Maximum heading level to include (1-9)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Ensure max_level is within valid range
        max_level = max(1, min(max_level, 9))
        
        doc = Document(filename)
        
        # Collect headings and their positions
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph style is a heading
            if paragraph.style and paragraph.style.name.startswith('Heading '):
                try:
                    # Extract heading level from style name
                    level = int(paragraph.style.name.split(' ')[1])
                    if level <= max_level:
                        headings.append({
                            'level': level,
                            'text': paragraph.text,
                            'position': i
                        })
                except (ValueError, IndexError):
                    # Skip if heading level can't be determined
                    pass
        
        if not headings:
            return f"No headings found in document {filename}. Table of contents not created."
        
        # Create a new document with the TOC
        toc_doc = Document()
        
        # Add title
        if title:
            toc_doc.add_heading(title, level=1)
        
        # Add TOC entries
        for heading in headings:
            # Indent based on level (using tab characters)
            indent = '    ' * (heading['level'] - 1)
            toc_doc.add_paragraph(f"{indent}{heading['text']}")
        
        # Add page break
        toc_doc.add_page_break()
        
        # Get content from original document
        for paragraph in doc.paragraphs:
            p = toc_doc.add_paragraph(paragraph.text)
            # Copy style if possible
            try:
                if paragraph.style:
                    p.style = paragraph.style.name
            except:
                pass
        
        # Copy tables
        for table in doc.tables:
            # Create a new table with the same dimensions
            new_table = toc_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_table.cell(i, j).text = paragraph.text
        
        # Save the new document with TOC
        toc_doc.save(filename)
        
        return f"Table of contents with {len(headings)} entries added to {filename}"
    except Exception as e:
        return f"Failed to add table of contents: {str(e)}"


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def delete_paragraph_range_tool(filename: str, start_index: int, end_index: int) -> str:
    """Delete a range of paragraphs from a document."""
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"

    return delete_paragraph_range(filename, start_index, end_index)


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)

        count, in_hyperlink = find_and_replace_text_detailed(doc, find_text, replace_text)

        if count > 0:
            doc.save(filename)
            base = f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
            if in_hyperlink:
                base += f" ({in_hyperlink} inside hyperlink display text)"
            return base
        hint = diagnose_outline_prefix_miss(doc, find_text)
        if hint:
            return f"No occurrences of '{find_text}' found. {hint}"
        return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

async def insert_header_near_text_tool(filename: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_header_near_text(filename, target_text, header_title, position, header_style, target_paragraph_index)

async def insert_numbered_list_near_text_tool(filename: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """Insert a bulleted or numbered list before or after the target paragraph. Specify by text or paragraph index."""
    return insert_numbered_list_near_text(filename, target_text, list_items, position, target_paragraph_index, bullet_type)

async def insert_line_or_paragraph_near_text_tool(filename: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None, copy_style_from_index: int = None) -> str:
    """Insert a new line or paragraph (with specified or matched style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_line_or_paragraph_near_text(filename, target_text, line_text, position, line_style, target_paragraph_index, copy_style_from_index)

async def replace_paragraph_block_below_header_tool(filename: str, header_text: str, new_paragraphs: list, detect_block_end_fn=None) -> str:
    """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
    return replace_paragraph_block_below_header(filename, header_text, new_paragraphs, detect_block_end_fn)

async def replace_block_between_manual_anchors_tool(filename: str, start_anchor_text: str, new_paragraphs: list, end_anchor_text: str = None, match_fn=None, new_paragraph_style: str = None) -> str:
    """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
    return replace_block_between_manual_anchors(filename, start_anchor_text, new_paragraphs, end_anchor_text, match_fn, new_paragraph_style)

async def replace_paragraph_text_tool(filename: str, paragraph_index: int, new_text: str,
                                      preserve_style: bool = True, parse_markdown: bool = False) -> str:
    """Replace text of a specific paragraph by index."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}."
    return replace_paragraph_text(filename, paragraph_index, new_text, preserve_style, parse_markdown)

async def replace_paragraph_range_tool(filename: str, start_index: int, end_index: int,
                                        new_paragraphs: list, style: str = None,
                                        preserve_style: bool = False) -> str:
    """Replace a range of paragraphs in a single operation."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}."
    return replace_paragraph_range(filename, start_index, end_index, new_paragraphs, style, preserve_style)
