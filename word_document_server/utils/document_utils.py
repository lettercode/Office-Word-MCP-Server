"""
Document utility functions for Word Document Server.
"""
import json
import unicodedata
import re
import logging
from typing import Dict, List, Any, Optional
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.run import Run

logger = logging.getLogger(__name__)

# Qualified tag names for XML element comparison
_W_P = qn('w:p')
_W_TBL = qn('w:tbl')


def _parse_markdown_runs(text: str) -> list:
    """Parse markdown-style inline formatting into a list of run specs.

    Supports:
    - ***bold italic*** (must be checked first - greedy)
    - **bold**
    - *italic*

    Returns:
        List of dicts: [{"text": "...", "bold": bool, "italic": bool}, ...]
    """
    if not text:
        return [{"text": "", "bold": False, "italic": False}]

    # Pattern matches ***bold italic***, **bold**, or *italic* (greedy, non-nested)
    # Order matters: check *** before ** before *
    pattern = r'(\*{3}(.+?)\*{3}|\*{2}(.+?)\*{2}|\*(.+?)\*)'

    runs = []
    last_end = 0

    for match in re.finditer(pattern, text):
        start = match.start()

        # Add plain text before this match
        if start > last_end:
            runs.append({
                "text": text[last_end:start],
                "bold": False,
                "italic": False
            })

        # Determine which group matched
        if match.group(2) is not None:
            # *** bold italic ***
            runs.append({
                "text": match.group(2),
                "bold": True,
                "italic": True
            })
        elif match.group(3) is not None:
            # ** bold **
            runs.append({
                "text": match.group(3),
                "bold": True,
                "italic": False
            })
        elif match.group(4) is not None:
            # * italic *
            runs.append({
                "text": match.group(4),
                "bold": False,
                "italic": True
            })

        last_end = match.end()

    # Add remaining text after last match
    if last_end < len(text):
        runs.append({
            "text": text[last_end:],
            "bold": False,
            "italic": False
        })

    # If no matches found, return the whole text as plain
    if not runs:
        runs.append({"text": text, "bold": False, "italic": False})

    return runs


def _normalize_text(s: str) -> str:
    """Normalize text for reliable matching: NFKC normalize, collapse whitespace, strip."""
    s = unicodedata.normalize("NFKC", s)
    s = re.sub(r'\s+', ' ', s)
    return s.strip()


def get_document_properties(doc_path: str, include_outline: bool = False) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)
        core_props = doc.core_properties

        result = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "word_count_method": "body_paragraphs_whitespace_split",
            "word_count_note": (
                "Counts words in body paragraphs only using whitespace splitting. "
                "Does not include text in tables, headers, footers, footnotes, "
                "or textboxes. May differ from Microsoft Word's built-in word count."
            ),
            "table_word_count": sum(
                len(para.text.split())
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
                for para in cell.paragraphs
            ),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }

        if include_outline:
            headings = []
            for i, para in enumerate(doc.paragraphs):
                if para.style and para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.split(" ")[1])
                    except (ValueError, IndexError):
                        level = 0
                    headings.append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name,
                        "level": level
                    })
            result["headings"] = headings

        return result
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs


_W_HYPERLINK = qn('w:hyperlink')


def iter_paragraph_runs_deep(para):
    """Yield ``Run`` objects wrapping every ``<w:r>`` descendant of the
    paragraph element in document order — including runs nested inside
    ``<w:hyperlink>``, ``<w:sdt>``, ``<w:smartTag>``, ``<w:fldSimple>``,
    ``<w:ins>``, ``<w:del>``. This is the right iterator for any tool that
    needs to see *user-visible* text, since ``Paragraph.runs`` only returns
    direct ``<w:r>`` children and silently skips hyperlink display text
    (Bug B)."""
    for r_elem in para._p.iter(qn('w:r')):
        yield Run(r_elem, para)


def paragraph_full_text(para):
    """Concatenated text of every ``<w:r>`` descendant of the paragraph,
    including hyperlink display text. Use this in preference to
    ``Paragraph.text`` when the search must cover hyperlinks reliably across
    python-docx versions."""
    return "".join(r.text for r in iter_paragraph_runs_deep(para))


def _run_is_inside_hyperlink(run_elem):
    """Return True if any ancestor of ``run_elem`` is a ``<w:hyperlink>``."""
    parent = run_elem.getparent()
    while parent is not None:
        if parent.tag == _W_HYPERLINK:
            return True
        parent = parent.getparent()
    return False


def _cleanup_empty_hyperlinks(para):
    """Remove ``<w:hyperlink>`` elements whose visible text is now empty.

    A search-and-replace operation that clears the entire display label of a
    hyperlink leaves a zombie element behind: the link still exists in the
    XML and the rels part, but its on-screen footprint is zero — it renders
    as a clickable nothing. This helper drops such elements (Bug B
    Symptom 2). The ``.rels`` entry is intentionally left in place; Word
    tolerates orphan rels and rewriting the rels part safely is fragile.
    """
    w_t = qn('w:t')
    for hyper in list(para._p.iter(_W_HYPERLINK)):
        text = "".join((t.text or "") for t in hyper.iter(w_t))
        if text == "":
            parent = hyper.getparent()
            if parent is not None:
                parent.remove(hyper)


def _replace_in_paragraph(para, old_text, new_text):
    """Replace ``old_text`` with ``new_text`` in a paragraph.

    Returns ``(count, in_hyperlink_count)`` where ``count`` is the total
    replacements made in this paragraph and ``in_hyperlink_count`` is the
    subset whose match started inside a ``<w:hyperlink>`` element.

    The match scan runs once over the paragraph's full run text up front;
    matches are then applied right-to-left so leftward offsets stay valid as
    run text is rewritten. This guarantees termination even when
    ``new_text`` contains ``old_text`` (which previously caused an infinite
    loop in the naive ``while old_text in para.text`` implementation —
    Bug A). The run iteration walks every ``<w:r>`` descendant including
    those nested inside hyperlinks (Bug B).
    """
    if not old_text:
        return 0, 0
    runs = list(iter_paragraph_runs_deep(para))
    if not runs:
        return 0, 0
    runs_text = "".join(r.text for r in runs)

    positions = []
    start_pos = 0
    while True:
        pos = runs_text.find(old_text, start_pos)
        if pos < 0:
            break
        positions.append(pos)
        start_pos = pos + len(old_text)
    if not positions:
        return 0, 0

    char_map = []  # (run_index, offset_in_run) keyed by position in runs_text
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_map.append((ri, ci))

    in_hyperlink = 0
    for pos in reversed(positions):
        end = pos + len(old_text)
        start_ri, start_ci = char_map[pos]
        end_ri, end_ci = char_map[end - 1]
        if _run_is_inside_hyperlink(runs[start_ri]._element):
            in_hyperlink += 1
        if start_ri == end_ri:
            run = runs[start_ri]
            run.text = run.text[:start_ci] + new_text + run.text[end_ci + 1:]
        else:
            runs[start_ri].text = runs[start_ri].text[:start_ci] + new_text
            for ri in range(start_ri + 1, end_ri):
                runs[ri].text = ""
            runs[end_ri].text = runs[end_ri].text[end_ci + 1:]

    _cleanup_empty_hyperlinks(para)
    return len(positions), in_hyperlink


_OUTLINE_PREFIX_RE = re.compile(r"^(#{1,6})\s+(.+)$")


def diagnose_outline_prefix_miss(doc, find_text: str) -> Optional[str]:
    """Return a hint when `find_text` looks like a markdown outline prefix
    (e.g. '## Performance') but no paragraph contains the literal hashes —
    and a heading-styled paragraph *does* carry the suffix as its text.

    Heading paragraphs in a .docx store only the text; the level is encoded
    in the style (``Heading 1``..``Heading 6``). Callers who copy outline
    rows from ``get_document_outline`` can paste back the ``## `` prefix
    and get a silent no-match. This helper surfaces that specific mismatch.

    Returns None when no heading paragraph matches the suffix.
    """
    m = _OUTLINE_PREFIX_RE.match(find_text)
    if not m:
        return None
    hashes, suffix = m.group(1), m.group(2).strip()
    if not suffix:
        return None
    for i, para in enumerate(doc.paragraphs):
        if para.text != suffix:
            continue
        style_name = para.style.name if para.style else ""
        if not style_name.startswith("Heading "):
            continue
        hint = (
            f"Hint: heading paragraphs store text without '#' prefix. "
            f"Found '{suffix}' at paragraph index {i} (style: {style_name}). "
            f"Retry with find_text='{suffix}' or use replace_paragraph_text."
        )
        try:
            heading_level = int(style_name.split()[-1])
        except (ValueError, IndexError):
            heading_level = None
        if heading_level is not None and heading_level != len(hashes):
            hint += (
                f" (note: outline hash count {len(hashes)} does not match {style_name})"
            )
        return hint
    return None


def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document, skipping Table of Contents (TOC) paragraphs.
    Handles text that spans multiple XML runs within a paragraph and reaches
    text nested inside ``<w:hyperlink>`` elements (Bug B).

    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with

    Returns:
        Either an int total count (legacy callers) — see
        :func:`find_and_replace_text_detailed` for the structured form. To
        preserve the historical contract, this function returns just the
        total. Callers needing the per-location breakdown should use the
        ``_detailed`` variant.
    """
    total, _ = find_and_replace_text_detailed(doc, old_text, new_text)
    return total


def find_and_replace_text_detailed(doc, old_text, new_text):
    """Same as :func:`find_and_replace_text` but returns
    ``(total, in_hyperlink)`` so callers can disclose how many of the
    matches lived inside hyperlink display text."""
    total = 0
    in_hyperlink_total = 0

    def _is_toc(para):
        return bool(para.style and para.style.name.startswith("TOC"))

    for para in doc.paragraphs:
        if _is_toc(para):
            continue
        if old_text in paragraph_full_text(para):
            n, h = _replace_in_paragraph(para, old_text, new_text)
            total += n
            in_hyperlink_total += h

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _is_toc(para):
                        continue
                    if old_text in paragraph_full_text(para):
                        n, h = _replace_in_paragraph(para, old_text, new_text)
                        total += n
                        in_hyperlink_total += h

    return total, in_hyperlink_total


def get_document_xml(doc_path: str) -> str:
    """Extract and return the raw XML structure of the Word document (word/document.xml)."""
    import os
    import zipfile
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        with zipfile.ZipFile(doc_path) as docx_zip:
            with docx_zip.open('word/document.xml') as xml_file:
                return xml_file.read().decode('utf-8')
    except Exception as e:
        return f"Failed to extract XML: {str(e)}"


def insert_header_near_text(doc_path: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search."""
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        new_para = doc.add_paragraph(header_title, style=header_style)
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} paragraph (index {anchor_index})."
        else:
            return f"Header '{header_title}' (style: {header_style}) inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert header: {str(e)}"


def _copy_run_formatting(source_run, target_run):
    """Copy character-level formatting from source run to target run."""
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb


def insert_line_or_paragraph_near_text(doc_path: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None, copy_style_from_index: int = None) -> str:
    """
    Insert a new line or paragraph (with specified or matched style) before or after the target paragraph.
    You can specify the target by text (first match) or by paragraph index.
    Skips paragraphs whose style name starts with 'TOC' if using text search.
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine style: use provided or match target
        style = line_style if line_style else para.style
        new_para = doc.add_paragraph(line_text, style=style)
        # Copy run formatting if requested
        if copy_style_from_index is not None:
            if 0 <= copy_style_from_index < len(doc.paragraphs):
                source_para = doc.paragraphs[copy_style_from_index]
                if source_para.runs and new_para.runs:
                    _copy_run_formatting(source_para.runs[0], new_para.runs[0])
        if position == 'before':
            para._element.addprevious(new_para._element)
        else:
            para._element.addnext(new_para._element)
        doc.save(doc_path)
        if anchor_index is not None:
            return f"Line/paragraph inserted {position} paragraph (index {anchor_index}) with style '{style}'."
        else:
            return f"Line/paragraph inserted {position} the target paragraph with style '{style}'."
    except Exception as e:
        return f"Failed to insert line/paragraph: {str(e)}"


def add_bullet_numbering(paragraph, num_id=1, level=0):
    """
    Add bullet/numbering XML to a paragraph.

    Args:
        paragraph: python-docx Paragraph object
        num_id: Numbering definition ID (1=bullets, 2=numbers, etc.)
        level: Indentation level (0=first level, 1=second level, etc.)

    Returns:
        The modified paragraph
    """
    # Get or create paragraph properties
    pPr = paragraph._element.get_or_add_pPr()

    # Remove existing numPr if any (to avoid duplicates)
    existing_numPr = pPr.find(qn('w:numPr'))
    if existing_numPr is not None:
        pPr.remove(existing_numPr)

    # Create numbering properties element
    numPr = OxmlElement('w:numPr')

    # Set indentation level
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)

    # Set numbering definition ID
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)

    # Add to paragraph properties
    pPr.append(numPr)

    return paragraph


def insert_numbered_list_near_text(doc_path: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """
    Insert a bulleted or numbered list before or after the target paragraph. Specify by text or paragraph index. Skips TOC paragraphs in text search.
    Args:
        doc_path: Path to the Word document
        target_text: Text to search for in paragraphs (optional if using index)
        list_items: List of strings, each as a list item
        position: 'before' or 'after' (default: 'after')
        target_paragraph_index: Optional paragraph index to use as anchor
        bullet_type: 'bullet' for bullets (•), 'number' for numbers (1,2,3) (default: 'bullet')
    Returns:
        Status message
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        para = None
        if target_paragraph_index is not None:
            if target_paragraph_index < 0 or target_paragraph_index >= len(doc.paragraphs):
                return f"Invalid target_paragraph_index: {target_paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            para = doc.paragraphs[target_paragraph_index]
            found = True
        else:
            for i, p in enumerate(doc.paragraphs):
                # Skip TOC paragraphs
                if p.style and p.style.name.lower().startswith("toc"):
                    continue
                if target_text and target_text in p.text:
                    para = p
                    found = True
                    break
        if not found or para is None:
            return f"Target paragraph not found (by index or text). (TOC paragraphs are skipped in text search)"
        # Save anchor index before insertion
        if target_paragraph_index is not None:
            anchor_index = target_paragraph_index
        else:
            anchor_index = None
            for i, p in enumerate(doc.paragraphs):
                if p is para:
                    anchor_index = i
                    break
        # Determine numbering ID based on bullet_type
        num_id = 1 if bullet_type == 'bullet' else 2

        # Use ListParagraph style for proper list formatting
        style_name = None
        for candidate in ['List Paragraph', 'ListParagraph', 'Normal']:
            try:
                _ = doc.styles[candidate]
                style_name = candidate
                break
            except KeyError:
                continue
        if not style_name:
            style_name = None  # fallback to default

        new_paras = []
        for item in (list_items or []):
            p = doc.add_paragraph(item, style=style_name)
            # Add bullet numbering XML - this is the fix!
            add_bullet_numbering(p, num_id=num_id, level=0)
            new_paras.append(p)
        # Move the new paragraphs to the correct position
        for p in reversed(new_paras):
            if position == 'before':
                para._element.addprevious(p._element)
            else:
                para._element.addnext(p._element)
        doc.save(doc_path)
        list_type = "bulleted" if bullet_type == 'bullet' else "numbered"
        if anchor_index is not None:
            return f"{list_type.capitalize()} list with {len(new_paras)} items inserted {position} paragraph (index {anchor_index})."
        else:
            return f"{list_type.capitalize()} list with {len(new_paras)} items inserted {position} the target paragraph."
    except Exception as e:
        return f"Failed to insert numbered list: {str(e)}"


def replace_paragraph_text(doc_path: str, paragraph_index: int, new_text: str,
                           preserve_style: bool = True, parse_markdown: bool = False) -> str:
    """Replace the text of a paragraph at a given index, optionally preserving style.

    Args:
        doc_path: Path to the Word document
        paragraph_index: Index of the paragraph to replace
        new_text: New text content
        preserve_style: Preserve paragraph-level style (default True)
        parse_markdown: Parse *italic*, **bold**, ***bold italic*** (default False)
    """
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"

    try:
        doc = Document(doc_path)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."

        para = doc.paragraphs[paragraph_index]
        old_style = para.style

        # Clear all existing runs by removing their XML elements
        for run in list(para.runs):
            run.text = ""
        for run in list(para.runs):
            if hasattr(run, '_r') and run._r.getparent() is not None:
                run._r.getparent().remove(run._r)

        if parse_markdown:
            # Parse markdown and create formatted runs
            run_specs = _parse_markdown_runs(new_text)
            for spec in run_specs:
                run = para.add_run(spec["text"])
                if spec["bold"]:
                    run.bold = True
                if spec["italic"]:
                    run.italic = True
        else:
            # Original behavior: single unformatted run
            para.add_run(new_text)

        if preserve_style and old_style:
            para.style = old_style
        elif not preserve_style:
            para.style = doc.styles["Normal"]

        doc.save(doc_path)
        return f"Paragraph at index {paragraph_index} replaced successfully."
    except Exception as e:
        return f"Failed to replace paragraph: {str(e)}"


def replace_paragraph_range(doc_path: str, start_index: int, end_index: int,
                            new_paragraphs: list, style: str = None,
                            preserve_style: bool = False) -> str:
    """Replace paragraphs from start_index to end_index (inclusive) with new_paragraphs.

    Args:
        doc_path: Path to the document
        start_index: First paragraph index to replace (inclusive)
        end_index: Last paragraph index to replace (inclusive)
        new_paragraphs: List of text strings for new paragraphs
        style: Optional style name for new paragraphs (overrides preserve_style)
        preserve_style: If True, copies style from the paragraph at start_index
    """
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"

    try:
        doc = Document(doc_path)
        total = len(doc.paragraphs)

        if start_index < 0 or end_index >= total or start_index > end_index:
            return f"Invalid range [{start_index}, {end_index}]. Document has {total} paragraphs (0-{total-1})."

        # Determine style for new paragraphs
        if style:
            style_to_use = style
        elif preserve_style:
            style_to_use = doc.paragraphs[start_index].style.name if doc.paragraphs[start_index].style else "Normal"
        else:
            style_to_use = "Normal"

        # Get anchor element (paragraph before start_index)
        if start_index > 0:
            anchor_element = doc.paragraphs[start_index - 1]._element
        else:
            anchor_element = None

        # Remove paragraphs in range (reverse to preserve indices)
        for i in range(end_index, start_index - 1, -1):
            p = doc.paragraphs[i]._p
            p.getparent().remove(p)

        # Insert new paragraphs
        body = doc.element.body

        prev_element = anchor_element
        for text in new_paragraphs:
            new_para = doc.add_paragraph(text, style=style_to_use)
            new_p_element = new_para._element
            # Remove from end of body (where add_paragraph appends it)
            body.remove(new_p_element)
            # Insert at correct position
            if prev_element is not None:
                prev_element.addnext(new_p_element)
            else:
                body.insert(0, new_p_element)
            prev_element = new_p_element

        doc.save(doc_path)
        removed = end_index - start_index + 1
        return f"Replaced {removed} paragraph(s) (indices {start_index}-{end_index}) with {len(new_paragraphs)} new paragraph(s)."
    except Exception as e:
        return f"Failed to replace paragraph range: {str(e)}"


def delete_paragraph_range(doc_path: str, start_index: int, end_index: int) -> str:
    """Delete paragraphs from start_index to end_index inclusive.

    Removes XML elements in reverse order to preserve internal indices.
    When making multiple range deletions, process higher indices first.

    Args:
        doc_path: Path to the Word document
        start_index: First paragraph index to delete (inclusive, 0-based)
        end_index: Last paragraph index to delete (inclusive, 0-based)

    Returns:
        Success or error message string
    """
    import os
    if not os.path.exists(doc_path):
        return f"Error: Document {doc_path} does not exist"

    try:
        doc = Document(doc_path)
        total = len(doc.paragraphs)

        if start_index < 0:
            return f"Error: start_index ({start_index}) must be >= 0"
        if end_index >= total:
            return f"Error: end_index ({end_index}) exceeds paragraph count ({total})"
        if start_index > end_index:
            return f"Error: start_index ({start_index}) > end_index ({end_index})"

        # Delete in reverse order to preserve indices
        for i in range(end_index, start_index - 1, -1):
            p = doc.paragraphs[i]._p
            p.getparent().remove(p)

        doc.save(doc_path)
        count = end_index - start_index + 1
        return f"Successfully deleted {count} paragraph(s) (indices {start_index}-{end_index})."
    except Exception as e:
        return f"Error: Failed to delete paragraph range: {str(e)}"


def is_toc_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de tabla de contenido (TOC)."""
    return para.style and para.style.name.upper().startswith("TOC")


def is_heading_paragraph(para):
    """Devuelve True si el párrafo tiene un estilo de encabezado (Heading 1, Heading 2, etc)."""
    return para.style and para.style.name.lower().startswith("heading")


# --- Helper: Get style name from a <w:p> element ---
def get_paragraph_style(el):
    from docx.oxml.ns import qn
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and 'w:val' in pStyle.attrib:
            return pStyle.attrib['w:val']
    return None

# --- Main: Delete everything under a header until next heading/TOC ---
def delete_block_under_header(doc, header_text):
    """
    Remove all elements (paragraphs, tables, etc.) after the header (by text) and before the next heading/TOC (by style).
    Returns: (header_element, elements_removed)
    """
    # Find the header paragraph by text — Pass 1: exact normalized match
    header_para = None
    header_idx = None
    normalized_header = _normalize_text(header_text).lower()

    for i, para in enumerate(doc.paragraphs):
        if _normalize_text(para.text).lower() == normalized_header:
            header_para = para
            header_idx = i
            break

    # Pass 2: contains match, prefer heading-styled paragraphs
    if header_para is None:
        for i, para in enumerate(doc.paragraphs):
            if is_heading_paragraph(para) and normalized_header in _normalize_text(para.text).lower():
                logger.info(f"Header matched via contains: '{para.text}' contains '{header_text}'")
                header_para = para
                header_idx = i
                break

    if header_para is None:
        return None, 0
    
    # Find the next heading/TOC paragraph to determine the end of the block
    end_idx = None
    for i in range(header_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style and para.style.name.lower().startswith(('heading', 'título', 'toc')):
            end_idx = i
            break
    
    # If no next heading found, delete until end of document
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    # Remove paragraphs by index (like delete_paragraph does)
    removed_count = 0
    for i in range(header_idx + 1, end_idx):
        if i < len(doc.paragraphs):  # Safety check
            para = doc.paragraphs[header_idx + 1]  # Always remove the first paragraph after header
            p = para._p
            p.getparent().remove(p)
            removed_count += 1
    
    return header_para._p, removed_count

# --- Usage in replace_paragraph_block_below_header ---
def replace_paragraph_block_below_header(
    doc_path: str,
    header_text: str,
    new_paragraphs: list,
    detect_block_end_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Reemplaza todo el contenido debajo de una cabecera (por texto), hasta el siguiente encabezado/TOC (por estilo).
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    
    doc = Document(doc_path)
    
    # Find the header paragraph first — Pass 1: exact normalized match
    header_para = None
    header_idx = None
    normalized_header = _normalize_text(header_text).lower()

    for i, para in enumerate(doc.paragraphs):
        if is_toc_paragraph(para):
            continue
        if _normalize_text(para.text).lower() == normalized_header:
            header_para = para
            header_idx = i
            break

    # Pass 2: contains match, prefer heading-styled paragraphs
    if header_para is None:
        for i, para in enumerate(doc.paragraphs):
            if is_toc_paragraph(para):
                continue
            if is_heading_paragraph(para) and normalized_header in _normalize_text(para.text).lower():
                logger.info(f"Header matched via contains: '{para.text}' contains '{header_text}'")
                header_para = para
                header_idx = i
                break

    if header_para is None:
        return f"Header '{header_text}' not found in document."
    
    # Delete everything under the header using the same document instance
    header_el, removed_count = delete_block_under_header(doc, header_text)
    
    # Now insert new paragraphs after the header (which should still be in the document)
    style_to_use = new_paragraph_style or "Normal"
    
    # Find the header again after deletion (it should still be there)
    current_para = header_para
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        current_para._element.addnext(new_para._element)
        current_para = new_para
    
    doc.save(doc_path)
    return f"Replaced content under '{header_text}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {removed_count} elements."


def replace_block_between_manual_anchors(
    doc_path: str,
    start_anchor_text: str,
    new_paragraphs: list,
    end_anchor_text: str = None,
    match_fn=None,
    new_paragraph_style: str = None
) -> str:
    """
    Replace all content (paragraphs, tables, etc.) between start_anchor_text and end_anchor_text (or next logical header if not provided).
    If end_anchor_text is None, deletes until next visually distinct paragraph (bold, all caps, or different font size), or end of document.
    Inserts new_paragraphs after the start anchor.
    """
    from docx import Document
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} not found."
    doc = Document(doc_path)
    body = doc.element.body
    elements = list(body)
    start_idx = None
    end_idx = None
    # Find start anchor — Pass 1: exact normalized match
    for i, el in enumerate(elements):
        if el.tag == _W_P:
            p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
            if match_fn:
                if match_fn(p_text, el):
                    start_idx = i
                    break
            elif _normalize_text(p_text) == _normalize_text(start_anchor_text):
                start_idx = i
                break
    # Pass 2: contains fallback
    if start_idx is None and not match_fn:
        normalized_start = _normalize_text(start_anchor_text)
        for i, el in enumerate(elements):
            if el.tag == _W_P:
                p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                if normalized_start in _normalize_text(p_text):
                    logger.info(f"Start anchor matched via contains: '{p_text}'")
                    start_idx = i
                    break
    if start_idx is None:
        return f"Start anchor '{start_anchor_text}' not found."
    # Find end anchor
    if end_anchor_text:
        # Pass 1: exact normalized match
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == _W_P:
                p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                if match_fn:
                    if match_fn(p_text, el, is_end=True):
                        end_idx = i
                        break
                elif _normalize_text(p_text) == _normalize_text(end_anchor_text):
                    end_idx = i
                    break
        # Pass 2: contains fallback for end anchor
        if end_idx is None and not match_fn:
            normalized_end = _normalize_text(end_anchor_text)
            for i in range(start_idx + 1, len(elements)):
                el = elements[i]
                if el.tag == _W_P:
                    p_text = "".join([node.text or '' for node in el.iter() if node.tag.endswith('}t')]).strip()
                    if normalized_end in _normalize_text(p_text):
                        logger.info(f"End anchor matched via contains: '{p_text}'")
                        end_idx = i
                        break
    else:
        # Heuristic: next visually distinct paragraph (bold, all caps, or different font size), or end of document
        for i in range(start_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == _W_P:
                # Check for bold, all caps, or font size
                runs = [node for node in el.iter() if node.tag.endswith('}r')]
                for run in runs:
                    rpr = run.find(qn('w:rPr'))
                    if rpr is not None:
                        if rpr.find(qn('w:b')) is not None or rpr.find(qn('w:caps')) is not None or rpr.find(qn('w:sz')) is not None:
                            end_idx = i
                            break
                if end_idx is not None:
                    break
    # Mark elements for removal
    to_remove = []
    for i in range(start_idx + 1, end_idx if end_idx is not None else len(elements)):
        to_remove.append(elements[i])
    for el in to_remove:
        body.remove(el)
    doc.save(doc_path)
    # Reload and find start anchor for insertion
    doc = Document(doc_path)
    paras = doc.paragraphs
    anchor_idx = None
    for i, para in enumerate(paras):
        if _normalize_text(para.text) == _normalize_text(start_anchor_text):
            anchor_idx = i
            break
    # Contains fallback for re-find
    if anchor_idx is None:
        normalized_start = _normalize_text(start_anchor_text)
        for i, para in enumerate(paras):
            if normalized_start in _normalize_text(para.text):
                anchor_idx = i
                break
    if anchor_idx is None:
        return f"Start anchor '{start_anchor_text}' not found after deletion (unexpected)."
    anchor_para = paras[anchor_idx]
    style_to_use = new_paragraph_style or "Normal"
    for text in new_paragraphs:
        new_para = doc.add_paragraph(text, style=style_to_use)
        anchor_para._element.addnext(new_para._element)
        anchor_para = new_para
    doc.save(doc_path)
    return f"Replaced content between '{start_anchor_text}' and '{end_anchor_text or 'next logical header'}' with {len(new_paragraphs)} paragraph(s), style: {style_to_use}, removed {len(to_remove)} elements."
