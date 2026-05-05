"""
Extended document utilities for Word Document Server.
"""
from typing import Dict, List, Any, Tuple
from docx import Document

from word_document_server.utils.document_utils import _normalize_text


def get_paragraph_text(doc_path: str, paragraph_index: int) -> Dict[str, Any]:
    """
    Get text from a specific paragraph in a Word document.
    
    Args:
        doc_path: Path to the Word document
        paragraph_index: Index of the paragraph to extract (0-based)
    
    Returns:
        Dictionary with paragraph text and metadata
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        
        # Check if paragraph index is valid
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {"error": f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."}
        
        paragraph = doc.paragraphs[paragraph_index]
        
        return {
            "index": paragraph_index,
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "is_heading": paragraph.style.name.startswith("Heading") if paragraph.style else False
        }
    except Exception as e:
        return {"error": f"Failed to get paragraph text: {str(e)}"}


def get_paragraph_range(doc_path: str, start_index: int, end_index: int) -> Dict[str, Any]:
    """Get text from a range of paragraphs (start to end index inclusive).

    Args:
        doc_path: Path to the Word document
        start_index: First paragraph index (inclusive, 0-based)
        end_index: Last paragraph index (inclusive, 0-based)

    Returns:
        Dict with "paragraphs" list (each having index, text, style, is_heading)
        and "count" field. Or dict with "error" key on failure.
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)
        total = len(doc.paragraphs)

        if start_index < 0:
            return {"error": f"start_index ({start_index}) must be >= 0"}
        if end_index >= total:
            return {"error": f"end_index ({end_index}) exceeds paragraph count ({total})"}
        if start_index > end_index:
            return {"error": f"start_index ({start_index}) > end_index ({end_index})"}

        paragraphs = []
        for i in range(start_index, end_index + 1):
            para = doc.paragraphs[i]
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "is_heading": para.style.name.startswith("Heading") if para.style else False
            })

        return {
            "paragraphs": paragraphs,
            "count": len(paragraphs)
        }
    except Exception as e:
        return {"error": f"Failed to get paragraph range: {str(e)}"}


from word_document_server.utils.document_utils import paragraph_full_text


def find_text(doc_path: str, text_to_find: str, match_case: bool = True, whole_word: bool = False, include_paragraph_text: bool = False) -> Dict[str, Any]:
    """
    Find all occurrences of specific text in a Word document.
    
    Args:
        doc_path: Path to the Word document
        text_to_find: Text to search for
        match_case: Whether to perform case-sensitive search
        whole_word: Whether to match whole words only
    
    Returns:
        Dictionary with search results
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    if not text_to_find:
        return {"error": "Search text cannot be empty"}
    
    try:
        doc = Document(doc_path)
        results = {
            "query": text_to_find,
            "match_case": match_case,
            "whole_word": whole_word,
            "occurrences": [],
            "total_count": 0
        }
        
        # Search in paragraphs
        for i, para in enumerate(doc.paragraphs):
            # Use deep-walk text so hyperlink display text is searchable
            # regardless of python-docx version (Bug B Symptom 3).
            full_text = paragraph_full_text(para)
            para_text = full_text
            search_text = text_to_find
            
            if not match_case:
                para_text = para_text.lower()
                search_text = search_text.lower()
            
            # Find all occurrences (simple implementation)
            start_pos = 0
            while True:
                if whole_word:
                    # For whole word search, we need to check word boundaries
                    words = para_text.split()
                    found = False
                    for word_idx, word in enumerate(words):
                        if (word == search_text or
                            (not match_case and word.lower() == search_text.lower())):
                            occurrence = {
                                "paragraph_index": i,
                                "position": word_idx,
                            }
                            if include_paragraph_text:
                                occurrence["text"] = full_text
                                occurrence["style"] = para.style.name if para.style else "Normal"
                            else:
                                occurrence["context"] = full_text[:100] + ("..." if len(full_text) > 100 else "")
                            results["occurrences"].append(occurrence)
                            results["total_count"] += 1
                            found = True

                    # Break after checking all words
                    break
                else:
                    # For substring search
                    pos = para_text.find(search_text, start_pos)
                    if pos == -1:
                        break

                    occurrence = {
                        "paragraph_index": i,
                        "position": pos,
                    }
                    if include_paragraph_text:
                        occurrence["text"] = full_text
                        occurrence["style"] = para.style.name if para.style else "Normal"
                    else:
                        occurrence["context"] = full_text[:100] + ("..." if len(full_text) > 100 else "")
                    results["occurrences"].append(occurrence)
                    results["total_count"] += 1
                    start_pos = pos + len(search_text)

        # Search in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        # Use deep-walk text so hyperlink display text in
                        # table cells is searchable too (Bug B Symptom 3).
                        full_text = paragraph_full_text(para)
                        para_text = full_text
                        search_text = text_to_find

                        if not match_case:
                            para_text = para_text.lower()
                            search_text = search_text.lower()

                        # Find all occurrences (simple implementation)
                        start_pos = 0
                        while True:
                            if whole_word:
                                # For whole word search, check word boundaries
                                words = para_text.split()
                                found = False
                                for word_idx, word in enumerate(words):
                                    if (word == search_text or
                                        (not match_case and word.lower() == search_text.lower())):
                                        occurrence = {
                                            "location": f"Table {table_idx}, Row {row_idx}, Column {col_idx}",
                                            "position": word_idx,
                                        }
                                        if include_paragraph_text:
                                            occurrence["text"] = full_text
                                            occurrence["style"] = para.style.name if para.style else "Normal"
                                        else:
                                            occurrence["context"] = full_text[:100] + ("..." if len(full_text) > 100 else "")
                                        results["occurrences"].append(occurrence)
                                        results["total_count"] += 1
                                        found = True

                                # Break after checking all words
                                break
                            else:
                                # For substring search
                                pos = para_text.find(search_text, start_pos)
                                if pos == -1:
                                    break

                                occurrence = {
                                    "location": f"Table {table_idx}, Row {row_idx}, Column {col_idx}",
                                    "position": pos,
                                }
                                if include_paragraph_text:
                                    occurrence["text"] = full_text
                                    occurrence["style"] = para.style.name if para.style else "Normal"
                                else:
                                    occurrence["context"] = full_text[:100] + ("..." if len(full_text) > 100 else "")
                                results["occurrences"].append(occurrence)
                                results["total_count"] += 1
                                start_pos = pos + len(search_text)
        
        return results
    except Exception as e:
        return {"error": f"Failed to search for text: {str(e)}"}


def get_section_paragraphs(doc_path: str, heading_text: str, include_heading: bool = True) -> Dict[str, Any]:
    """Get all paragraphs under a heading until the next same-or-higher-level heading.

    Uses normalized text matching (NFKC + whitespace collapse) to find the heading.
    Falls back to substring matching if exact match fails.

    Args:
        doc_path: Path to the Word document
        heading_text: Text of the heading to find
        include_heading: Whether to include the heading paragraph itself (default True)

    Returns:
        Dict with heading metadata and paragraphs list.
        Or dict with "error" key on failure.
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    try:
        doc = Document(doc_path)
        normalized_search = _normalize_text(heading_text)

        # Find the heading paragraph — exact normalized match first
        heading_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name.startswith("Heading"):
                normalized_para = _normalize_text(para.text)
                if normalized_para == normalized_search:
                    heading_idx = i
                    break

        # Fallback: substring match on heading paragraphs
        if heading_idx is None:
            for i, para in enumerate(doc.paragraphs):
                if para.style and para.style.name.startswith("Heading"):
                    if normalized_search in _normalize_text(para.text):
                        heading_idx = i
                        break

        if heading_idx is None:
            return {"error": f"Heading '{heading_text}' not found in document"}

        heading_para = doc.paragraphs[heading_idx]
        heading_style = heading_para.style.name if heading_para.style else "Heading 1"

        # Extract heading level number
        try:
            heading_level = int(heading_style.split(" ")[1])
        except (ValueError, IndexError):
            heading_level = 1

        # Walk forward collecting paragraphs until next same-or-higher-level heading
        next_heading_idx = None

        for i in range(heading_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if para.style and para.style.name.startswith("Heading"):
                try:
                    para_level = int(para.style.name.split(" ")[1])
                except (ValueError, IndexError):
                    para_level = 1
                if para_level <= heading_level:
                    next_heading_idx = i
                    break

        # Determine content boundaries
        if next_heading_idx is None:
            content_end_idx = len(doc.paragraphs) - 1
        else:
            content_end_idx = next_heading_idx - 1

        # Build paragraphs list
        paragraphs = []
        start = heading_idx if include_heading else heading_idx + 1
        end = content_end_idx

        for i in range(start, end + 1):
            para = doc.paragraphs[i]
            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "is_heading": para.style.name.startswith("Heading") if para.style else False
            })

        return {
            "heading_index": heading_idx,
            "heading_text": heading_para.text,
            "heading_style": heading_style,
            "heading_level": heading_level,
            "content_start_index": heading_idx + 1 if heading_idx + 1 <= content_end_idx else None,
            "content_end_index": content_end_idx if content_end_idx > heading_idx else None,
            "next_heading_index": next_heading_idx,
            "paragraphs": paragraphs
        }
    except Exception as e:
        return {"error": f"Failed to get section paragraphs: {str(e)}"}


def find_texts(doc_path: str, texts_to_find: List[str], match_case: bool = True,
               include_paragraph_text: bool = False) -> Dict[str, Any]:
    """Find multiple text strings in a document in a single pass.

    Loads the document once and checks each paragraph against all search strings.

    Args:
        doc_path: Path to the Word document
        texts_to_find: List of strings to search for
        match_case: Case-sensitive matching (default True)
        include_paragraph_text: Include full paragraph text in results (default False)

    Returns:
        Dict keyed by search string. Each value has "occurrences" list and "total_count",
        matching the shape of find_text() results.
        Or dict with "error" key on failure.
    """
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}

    if not texts_to_find:
        return {}

    try:
        doc = Document(doc_path)

        # Initialize results for each unique search string
        unique_texts = list(dict.fromkeys(texts_to_find))  # preserve order, dedupe
        results = {}
        for text in unique_texts:
            results[text] = {
                "occurrences": [],
                "total_count": 0
            }

        # Single pass through all paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text

            for search_text in unique_texts:
                compare_para = para_text if match_case else para_text.lower()
                compare_search = search_text if match_case else search_text.lower()

                start_pos = 0
                while True:
                    pos = compare_para.find(compare_search, start_pos)
                    if pos == -1:
                        break

                    occurrence = {
                        "paragraph_index": i,
                        "position": pos,
                    }
                    if include_paragraph_text:
                        occurrence["text"] = para.text
                        occurrence["style"] = para.style.name if para.style else "Normal"
                    else:
                        occurrence["context"] = para.text[:100] + ("..." if len(para.text) > 100 else "")

                    results[search_text]["occurrences"].append(occurrence)
                    results[search_text]["total_count"] += 1
                    start_pos = pos + len(compare_search)

        return results
    except Exception as e:
        return {"error": f"Failed to search for texts: {str(e)}"}
