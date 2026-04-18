"""Run-level normalization for anchor-based document edits.

The docx-editor library matches anchor text within a single `<w:r>` run
whereas python-docx-level helpers (``find_text_in_document``) search on
paragraph text, which spans runs. The gap produces silent "anchor not
found" errors in ``add_comment`` and the ``*_with_track_changes`` tools
whenever the anchor crosses a run boundary (e.g. a style change, a spell-
check rsid split, or an rsid boundary from a previous edit).

``normalize_paragraph_runs_for_anchor`` looks for ``anchor_text`` in each
paragraph's concatenated text; when the anchor is present but split across
runs, it splits the covering runs so the anchor lives in exactly one run,
preserving the ``rPr`` of every source run across the split. Subsequent
docx-editor calls then resolve the anchor.
"""
from __future__ import annotations

import copy
from typing import List, Tuple

from docx import Document
from docx.oxml.ns import qn


def normalize_paragraph_runs_for_anchor(docx_path: str, anchor_text: str) -> bool:
    """Ensure ``anchor_text`` lives in a single run wherever it occurs.

    Returns True if the anchor is present in the document after normalization
    (either it already fit in one run, or it was successfully split-and-merged
    into one). Returns False if no paragraph text contains the anchor at all,
    so the caller can short-circuit to a "not found" error without rewriting
    the file.
    """
    if not anchor_text:
        return False

    doc = Document(docx_path)
    changed = False
    found_anywhere = False

    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text
        if anchor_text not in text:
            continue
        found_anywhere = True
        start = text.find(anchor_text)
        end = start + len(anchor_text)
        if _anchor_already_in_single_run(paragraph, start, end, anchor_text):
            continue
        if _split_runs_for_anchor(paragraph, start, end):
            changed = True

    if changed:
        doc.save(docx_path)
    return found_anywhere


# ── internals ────────────────────────────────────────────────────────────


def _iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _run_text_segments(paragraph) -> List[Tuple[object, str, int, int]]:
    """Return (run_element, text, start_char, end_char) for each run in the
    paragraph in document order. ``text`` reflects the concatenation of all
    ``<w:t>`` children of the run (no tab/br handling — matches how
    ``paragraph.text`` builds its result enough to keep offsets aligned for
    common anchors)."""
    segments: List[Tuple[object, str, int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        t = run.text or ""
        segments.append((run._element, t, cursor, cursor + len(t)))
        cursor += len(t)
    return segments


def _anchor_already_in_single_run(paragraph, start: int, end: int, anchor_text: str) -> bool:
    for run in paragraph.runs:
        if anchor_text in (run.text or ""):
            return True
    # Also true if the paragraph has exactly one run covering the span.
    segments = _run_text_segments(paragraph)
    for _, _, s, e in segments:
        if s <= start and end <= e:
            return True
    return False


def _split_runs_for_anchor(paragraph, start: int, end: int) -> bool:
    """Split runs so that the span [start, end) is covered by exactly one
    run. Preserves each source run's ``w:rPr`` on all split fragments.
    Returns True if any split was performed."""
    segments = _run_text_segments(paragraph)
    if not segments:
        return False

    # Find run(s) that cover start and end.
    did_split = False
    # First: split the run that contains `start` so that `start` is at a
    # run boundary.
    did_split |= _split_run_at_offset(paragraph, start)
    # Re-read after potential split.
    did_split |= _split_run_at_offset(paragraph, end)
    # After the two splits, the anchor text may span multiple runs with the
    # same (or compatible) rPr — merge consecutive runs that together form
    # exactly the anchor. Simplest: do nothing else. docx-editor only
    # requires the anchor characters to live in a contiguous subsequence of
    # runs? No — it requires a single run. So we must merge.
    did_split |= _merge_runs_covering_anchor(paragraph, start, end)
    return did_split


def _split_run_at_offset(paragraph, offset: int) -> bool:
    """If ``offset`` falls strictly inside a run's text, split that run into
    two siblings (before/after). Returns True if a split was performed."""
    segments = _run_text_segments(paragraph)
    for run_elem, text, s, e in segments:
        if s < offset < e:
            rel = offset - s
            before_text = text[:rel]
            after_text = text[rel:]
            _replace_run_text(run_elem, before_text)
            clone = copy.deepcopy(run_elem)
            _replace_run_text(clone, after_text)
            run_elem.addnext(clone)
            return True
    return False


def _merge_runs_covering_anchor(paragraph, start: int, end: int) -> bool:
    """After boundary splits, the anchor spans one or more adjacent runs;
    merge them into a single run carrying the ``rPr`` of the first covered
    run."""
    segments = _run_text_segments(paragraph)
    covered = [
        (run_elem, text)
        for run_elem, text, s, e in segments
        if start <= s and e <= end
    ]
    if len(covered) <= 1:
        return False

    first_elem, first_text = covered[0]
    merged_text = "".join(t for _, t in covered)
    _replace_run_text(first_elem, merged_text)
    for run_elem, _ in covered[1:]:
        run_elem.getparent().remove(run_elem)
    return True


def _replace_run_text(run_elem, new_text: str) -> None:
    """Replace all ``<w:t>`` children of the run with a single ``<w:t
    xml:space='preserve'>`` carrying ``new_text``. ``<w:rPr>`` and other
    siblings (e.g. ``<w:br>``) are left alone unless they are ``<w:t>``."""
    from lxml import etree

    w_t_tag = qn("w:t")
    for child in list(run_elem):
        if child.tag == w_t_tag:
            run_elem.remove(child)

    new_t = etree.SubElement(run_elem, w_t_tag)
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = new_text
